# -*- coding: utf-8 -*-
"""Native python-docx renderer that walks the lease HTML (same source as the PDF) into a clean, editable Word file."""
import re
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0b, 0x1f, 0x3a)
GRAY = RGBColor(0x66, 0x70, 0x85)
DKGRAY = RGBColor(0x33, 0x40, 0x4f)

def _classes(el):
    return el.get("class", []) if isinstance(el, Tag) else []

def _u(n):
    return "_" * max(4, int(n))

def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def _set_table_borders(table, color="b9c2d0", sz="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)

def _no_borders(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'), 'nil'); borders.append(e)
    tblPr.append(borders)

def _add_field(paragraph, field, size=7.5, color=GRAY):
    run = paragraph.add_run(); run.font.size = Pt(size); run.font.color.rgb = color
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = f' {field} '
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(i); run._r.append(e)

def _fill_width(el):
    style = el.get("style", "") if isinstance(el, Tag) else ""
    m = re.search(r'min-width:(\d+)pt', style)
    w = int(m.group(1)) if m else 90
    return max(6, int(w / 3.2))

def _emit_inline(par, node):
    """Append inline content of a node to a docx paragraph."""
    for child in node.children:
        if isinstance(child, NavigableString):
            txt = str(child).replace('\xa0', ' ')
            txt = re.sub(r'\s+', ' ', txt)
            if txt:
                par.add_run(txt)
        elif isinstance(child, Tag):
            cls = _classes(child)
            if child.name == 'br':
                par.add_run().add_break()
            elif child.name in ('strong', 'b'):
                r = par.add_run(child.get_text()); r.bold = True
            elif child.name in ('em', 'i'):
                r = par.add_run(child.get_text()); r.italic = True
            elif child.name == 'span' and 'fill' in cls:
                r = par.add_run(' ' + _u(_fill_width(child)) + ' ')
            elif child.name == 'span' and 'label-sm' in cls:
                r = par.add_run(' ' + child.get_text()); r.font.size = Pt(7.5); r.font.color.rgb = GRAY
            elif child.name == 'span':
                _emit_inline(par, child)
            elif child.name == 'div' and 'fill-line' in cls:
                par.add_run().add_break()
                par.add_run(_u(70))
            elif child.name == 'div' and 'label-sm' in cls:
                par.add_run().add_break()
                r = par.add_run(child.get_text()); r.font.size = Pt(7.5); r.font.color.rgb = GRAY
            elif child.name == 'div' and ('big-space' in cls):
                par.add_run().add_break()
            else:
                _emit_inline(par, child)

def _para(doc, node, style_size=10.5, justify=True, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _emit_inline(p, node)
    for r in p.runs:
        if not r.font.size:
            r.font.size = Pt(style_size)
    return p

def _grid_table(doc, table_el):
    rows = table_el.find_all('tr')
    header_cells = rows[0].find_all(['th', 'td'])
    ncol = len(header_cells)
    t = doc.add_table(rows=0, cols=ncol)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(t)
    # header
    hr = t.add_row().cells
    for j, hc in enumerate(header_cells):
        _set_cell_bg(hr[j].__wrapped__ if hasattr(hr[j], '__wrapped__') else hr[j], "0b1f3a")
        cell = hr[j]
        cell.text = ''
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        run = p.add_run(hc.get_text().replace('\xa0', ' ').strip())
        run.bold = True; run.font.size = Pt(8.6); run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    # body
    for row in rows[1:]:
        cells = row.find_all(['td', 'th'])
        rc = t.add_row().cells
        for j in range(ncol):
            txt = cells[j].get_text().replace('\xa0', ' ').strip() if j < len(cells) else ''
            rc[j].text = ''
            p = rc[j].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            run = p.add_run(txt); run.font.size = Pt(9.3)
            # give empty fill cells some height
            if not txt:
                p.paragraph_format.space_after = Pt(8)
    return t

def _sigrow(doc, sigrow_el):
    cols = sigrow_el.find_all('div', class_='sigcol', recursive=False)
    t = doc.add_table(rows=1, cols=len(cols))
    _no_borders(t)
    for j, col in enumerate(cols):
        cell = t.rows[0].cells[j]
        cell.text = ''
        first = True
        for sub in col.find_all('div', recursive=False):
            scls = _classes(sub)
            if 'sigline' in scls:
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                first = False
                p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(0)
                p.add_run(_u(46))
            elif 'sigcap' in scls:
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(sub.get_text().strip()); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x40,0x50,0x6a)
            elif 'ibox' in scls:
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                first = False
                line = sub.find('span', class_='line'); cap = sub.find('span', class_='cap')
                p.add_run(_u(30))
                if cap:
                    p.add_run().add_break()
                    r = p.add_run(cap.get_text().strip()); r.font.size = Pt(7); r.font.color.rgb = GRAY
    return t

def _initials_clause(doc, table_el):
    body = table_el.find('td', class_='body')
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    if body:
        for sub in body.children:
            if isinstance(sub, Tag) and sub.name == 'p':
                _emit_inline(p, sub)
            elif isinstance(sub, NavigableString):
                t = re.sub(r'\s+', ' ', str(sub))
                if t.strip(): p.add_run(t)
    for r in p.runs:
        if not r.font.size: r.font.size = Pt(10)
    ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ip.paragraph_format.space_after = Pt(6)
    r = ip.add_run("Landlord initials " + _u(10) + "        Tenant initials " + _u(10))
    r.font.size = Pt(8.5); r.font.color.rgb = GRAY

def _numbered_lines(doc, nl_el):
    for nl in nl_el.find_all('div', class_='nl'):
        n = nl.find('span', class_='n')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run((n.get_text().strip() if n else '') + '  ' + _u(85))
        r.font.size = Pt(10.5)

def _heading(doc, text, size, color, before=12, after=3, bold=True, align=None, border=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    if align is not None: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
    r.font.name = 'Arial'
    if border:
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'd6dce6')
        pbdr.append(bottom); pPr.append(pbdr)
    return p

def _note(doc, node, strong=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.12)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8'); left.set(qn('w:color'), 'b9772f' if strong else '7089b0')
    pbdr.append(left); pPr.append(pbdr)
    _emit_inline(p, node)
    for r in p.runs:
        r.font.size = Pt(9.4);
        if not r.font.color or r.font.color.rgb is None: r.font.color.rgb = DKGRAY

def walk_body(doc, body_html):
    """Walk the top-level flow of a document body HTML into the docx."""
    soup = BeautifulSoup(body_html, 'html.parser')
    for el in soup.children:
        if isinstance(el, NavigableString):
            continue
        if not isinstance(el, Tag):
            continue
        cls = _classes(el)
        if el.name == 'div' and 'doc-title' in cls:
            _heading(doc, el.get_text(' ', strip=True), 16, NAVY, before=2, after=2,
                     align=WD_ALIGN_PARAGRAPH.CENTER)
        elif el.name == 'div' and 'doc-subtitle' in cls:
            _heading(doc, el.get_text(' ', strip=True), 10, RGBColor(0x40,0x50,0x6a), before=0, after=4,
                     bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif el.name == 'div' and ('badge' in cls or (el.get('style') and 'text-align:center' in el.get('style',''))):
            badge = el.find('span', class_='badge')
            if badge:
                _heading(doc, badge.get_text(strip=True), 8, RGBColor(0x5a,0x3a,0x7a), before=2, after=2,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
        elif el.name == 'hr':
            hp = doc.add_paragraph(); hp.paragraph_format.space_after = Pt(4)
            pPr = hp._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'1')
            bottom.set(qn('w:color'),'0b1f3a'); pbdr.append(bottom); pPr.append(pbdr)
        elif el.name == 'p' and 'caption' in cls:
            p = _para(doc, el, style_size=9, justify=False, space_after=8)
            for r in p.runs: r.font.color.rgb = RGBColor(0x40,0x50,0x6a)
        elif el.name == 'p':
            _para(doc, el, space_after=3 if 'tight' in cls else 6)
        elif el.name == 'h2':
            _heading(doc, el.get_text(strip=True), 10.8, NAVY, border=True)
        elif el.name == 'h3':
            _heading(doc, el.get_text(strip=True), 10, RGBColor(0x22,0x31,0x4f), before=8, after=2)
        elif el.name == 'table' and 'grid' in cls:
            _grid_table(doc, el)
        elif el.name == 'table' and 'initials' in cls:
            _initials_clause(doc, el)
        elif el.name == 'div' and 'note' in cls:
            _note(doc, el, strong=('note-strong' in cls))
        elif el.name == 'div' and 'numbered-lines' in cls:
            _numbered_lines(doc, el)
        elif el.name == 'div' and 'sigrow' in cls:
            _sigrow(doc, el)
        elif el.name == 'div' and 'section-block' in cls:
            walk_body(doc, el.decode_contents())
        elif el.name == 'div' and 'keepwith' in cls:
            walk_body(doc, el.decode_contents())
        elif el.name == 'div':
            # generic container: recurse
            inner = el.decode_contents().strip()
            if inner:
                walk_body(doc, inner)

def new_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6); style.paragraph_format.line_spacing = 1.15
    sec = doc.sections[0]
    sec.page_height = Inches(11); sec.page_width = Inches(8.5)
    sec.top_margin = Inches(0.85); sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
    # footer
    footer = sec.footer
    fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("1658 4th Street, Coachella, California 92236   |   Page ")
    r.font.size = Pt(7.5); r.font.color.rgb = GRAY
    _add_field(fp, "PAGE"); r2 = fp.add_run(" of "); r2.font.size = Pt(7.5); r2.font.color.rgb = GRAY
    _add_field(fp, "NUMPAGES")
    return doc

def page_break(doc):
    from docx.enum.text import WD_BREAK
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def build_lease_docx(sections_html_list, out_path):
    """sections_html_list: list of body-HTML strings; page break inserted between them."""
    doc = new_document()
    for i, body in enumerate(sections_html_list):
        if i > 0:
            page_break(doc)
        walk_body(doc, body)
    doc.save(out_path)
    return out_path
