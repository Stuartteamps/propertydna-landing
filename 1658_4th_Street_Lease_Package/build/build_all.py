# -*- coding: utf-8 -*-
"""Render the full 1658 4th Street lease package: HTML -> PDF (WeasyPrint), DOCX, merged PDF."""
import os, subprocess, sys
from weasyprint import HTML
from pypdf import PdfReader, PdfWriter

sys.path.insert(0, os.path.dirname(__file__))
from styles import full_html
from lease import lease_body
import annexes as A

OUT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
BUILD = os.path.dirname(__file__)

FOOTER_LEASE = "Month-to-Month Residential Rental Agreement"
FOOTER_PKG = "Month-to-Month Lease Package"

def render_pdf(html_str, pdf_path):
    HTML(string=html_str).write_pdf(pdf_path)
    return pdf_path

def npages(pdf_path):
    return len(PdfReader(pdf_path).pages)

def build(sources_note, assumptions_html):
    # ---- assemble body sections ----
    lease_html = lease_body()
    house_html = A.house_rules_body()
    addterms_html = A.additional_terms_body()
    movein_html = A.movein_body()
    key_html = A.key_receipt_body()
    contact_html = A.contact_forms_body()
    receipt_html = A.rent_receipt_body()
    disclosure_html = A.disclosure_checklist_body(sources_note)
    exemption_html = A.optional_exemption_body()
    guidance_html = A.owner_guidance_body(assumptions_html)

    # ---- LEASE (core + house rules + additional terms) : PDF + source HTML for DOCX ----
    lease_full = (lease_html
                  + '<div class="pagebreak"></div>' + house_html
                  + '<div class="pagebreak"></div>' + addterms_html)
    lease_doc = full_html("1658 4th Street Month-to-Month Lease", lease_full, FOOTER_LEASE)
    lease_html_path = os.path.join(BUILD, "_lease_source.html")
    with open(lease_html_path, "w") as f:
        f.write(lease_doc)
    lease_pdf = render_pdf(lease_doc, os.path.join(OUT, "1658_4th_Street_Month_to_Month_Lease.pdf"))

    # ---- Owner guidance PDF ----
    guidance_pdf = render_pdf(full_html("Owner Guidance", guidance_html, "Owner Guidance — Not Part of the Lease"),
                              os.path.join(OUT, "1658_4th_Street_Owner_Guidance.pdf"))

    # ---- Disclosure checklist PDF (+ optional exemption addendum) ----
    disc_full = disclosure_html + '<div class="pagebreak"></div>' + exemption_html
    disc_pdf = render_pdf(full_html("Disclosure Checklist", disc_full, "Disclosure Worksheet"),
                          os.path.join(OUT, "1658_4th_Street_Disclosure_Checklist.pdf"))

    # ---- Intermediate PDFs for the remaining annexes (for the merged package) ----
    def annex_pdf(name, body, footer):
        p = os.path.join(BUILD, f"_{name}.pdf")
        return render_pdf(full_html(name, body, footer), p)

    p_house = annex_pdf("house_rules", house_html, "House Rules and Material Terms")
    p_addterms = annex_pdf("additional_terms", addterms_html, "Additional Agreed Terms")
    p_movein = annex_pdf("movein", movein_html, "Move-In Condition Checklist")
    p_key = annex_pdf("key_receipt", key_html, "Key and Access-Device Receipt")
    p_contact = annex_pdf("contact_forms", contact_html, "Tenant & Emergency Contact Forms")
    p_receipt = annex_pdf("rent_receipt", receipt_html, "Rent-Payment Receipt")

    # lease-core-only PDF (without house/addterms) for page counting
    lease_core_pdf = annex_pdf("lease_core", lease_html, FOOTER_LEASE)

    # ---- Complete package: ONE combined render for continuous "Page X of Y" numbering ----
    pb = '<div class="pagebreak"></div>'
    combined_body = pb.join([
        lease_html, house_html, addterms_html, movein_html, key_html,
        contact_html, receipt_html, disclosure_html, exemption_html, guidance_html,
    ])
    complete_pdf = render_pdf(
        full_html("1658 4th Street — Complete Lease Package", combined_body, "Complete Lease Package"),
        os.path.join(OUT, "1658_4th_Street_Complete_Lease_Package.pdf"))

    # ---- DOCX of the lease (native python-docx renderer, walks the same HTML source) ----
    from render_docx import build_lease_docx
    docx_out = os.path.join(OUT, "1658_4th_Street_Month_to_Month_Lease.docx")
    build_lease_docx([lease_html, house_html, addterms_html], docx_out)
    print("DOCX saved:", docx_out)

    # ---- report page counts ----
    print("PAGE COUNTS:")
    print("  Lease (core+rules+terms):", npages(lease_pdf))
    print("  Lease core only:", npages(lease_core_pdf))
    print("  Owner guidance:", npages(guidance_pdf))
    print("  Disclosure+exemption:", npages(disc_pdf))
    print("  COMPLETE PACKAGE:", npages(complete_pdf))
    return complete_pdf

def build_docx(lease_html_path):
    docx_out = os.path.join(OUT, "1658_4th_Street_Month_to_Month_Lease.docx")
    # LibreOffice converts HTML -> DOCX
    subprocess.run(["soffice", "--headless", "--convert-to", "docx", "--outdir", BUILD, lease_html_path],
                   check=True, capture_output=True, timeout=120)
    converted = os.path.join(BUILD, "_lease_source.docx")
    # post-process: page size Letter, margins, footer with address + page number
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(converted)
    for sec in doc.sections:
        sec.page_height = Inches(11)
        sec.page_width = Inches(8.5)
        sec.top_margin = Inches(0.85); sec.bottom_margin = Inches(0.9)
        sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
        footer = sec.footer
        footer.is_linked_to_previous = False
        # clear existing footer paragraphs
        for p in list(footer.paragraphs):
            p.text = ""
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("1658 4th Street, Coachella, California 92236   |   Page ")
        run.font.size = Pt(7.5); run.font.color.rgb = RGBColor(0x66, 0x70, 0x85)
        # PAGE field
        _add_field(fp, "PAGE", 7.5)
        run2 = fp.add_run(" of ")
        run2.font.size = Pt(7.5); run2.font.color.rgb = RGBColor(0x66, 0x70, 0x85)
        _add_field(fp, "NUMPAGES", 7.5)
    doc.save(docx_out)
    print("DOCX saved:", docx_out)

def _add_field(paragraph, field, size_pt):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run()
    run.font.size = Pt(size_pt); run.font.color.rgb = RGBColor(0x66, 0x70, 0x85)
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = f' {field} '
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldEnd)

if __name__ == "__main__":
    # placeholders; overwritten by finalize step
    build("Sources verified against official California legislative materials; see California_Law_Sources.txt.",
          "<p>See California_Law_Sources.txt.</p>")
